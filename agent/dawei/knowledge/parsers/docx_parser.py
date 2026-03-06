# Copyright (c) 2025 格律至微
# SPDX-License-Identifier: AGPL-3.0-only

"""DOCX document parser"""

import hashlib
from pathlib import Path

from dawei.knowledge.models import Document, DocumentMetadata, DocumentType
from dawei.knowledge.parsers.base import BaseParser


class DocxParser(BaseParser):
    """DOCX document parser"""

    async def parse(self, file_path: str | Path) -> Document:
        """Parse DOCX document"""
        file_path = Path(file_path)

        # Calculate SHA256
        sha256_hash = self._calculate_sha256(file_path)

        # Extract DOCX content
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        doc = DocxDocument(file_path)
        content_parts = []

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading 1"):
                content_parts.append(f"# {para.text}")
            elif para.style.name.startswith("Heading 2"):
                content_parts.append(f"## {para.text}")
            elif para.style.name.startswith("Heading 3"):
                content_parts.append(f"### {para.text}")
            else:
                content_parts.append(para.text)

        full_content = "\n\n".join(content_parts)

        # Create metadata
        metadata = DocumentMetadata(
            file_path=str(file_path),
            file_name=file_path.name,
            file_size=file_path.stat().st_size,
            file_type=DocumentType.DOCX,
            sha256=sha256_hash,
        )

        return Document(id=sha256_hash, metadata=metadata, content=full_content)

    def supports_file_type(self, file_path: str | Path) -> bool:
        """Check if file is a DOCX"""
        return Path(file_path).suffix.lower() == ".docx"

    def _calculate_sha256(self, file_path: Path) -> str:
        """Calculate SHA256 hash of file"""
        sha256 = hashlib.sha256()
        with file_path.open("rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256.update(chunk)
        return sha256.hexdigest()
