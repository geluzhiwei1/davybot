# Copyright (c) 2025 格律至微
# SPDX-License-Identifier: AGPL-3.0-only

import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from pydantic import BaseModel, Field

from dawei.core.decorators import safe_tool_operation
from dawei.core.path_security import (
    PathTraversalError,
    safe_path_join,
    sanitize_filename,
)
from dawei.tools.custom_base_tool import CustomBaseTool

logger = logging.getLogger(__name__)


# Insert Content Tool
class InsertContentInput(BaseModel):
    """Input for InsertContentTool."""

    path: str = Field(..., description="File path to insert content into.")
    line: int = Field(
        ...,
        description="Line number where content will be inserted (1-based, 0 to append).",
    )
    content: str = Field(..., description="Content to insert at the specified line.")


class InsertContentTool(CustomBaseTool):
    """Tool for inserting content into files without modifying existing content."""

    name: str = "insert_content"
    description: str = "Adds new lines of content into a file without modifying existing content."
    args_schema: type[BaseModel] = InsertContentInput

    @safe_tool_operation("insert_content", fallback_value="Error: Failed to insert content")
    def _run(self, path: str, line: int, content: str) -> str:
        """Insert content at specified line."""
        # Get workspace directory from context
        workspace_dir = None
        if hasattr(self, "context") and self.context is not None:
            workspace_dir = getattr(self.context, "cwd", None)

        # Sanitize filename to prevent directory traversal
        sanitized_filename = sanitize_filename(Path(path).name)

        # Reconstruct path with sanitized filename
        path_obj = Path(path)
        dir_name = path_obj.parent
        path = str(dir_name / sanitized_filename) if dir_name != Path() else sanitized_filename

        # If we have a workspace directory, use safe path joining
        if workspace_dir:
            try:
                # Allow common file extensions
                allowed_extensions: set[str] = {
                    ".py",
                    ".js",
                    ".ts",
                    ".tsx",
                    ".jsx",
                    ".vue",
                    ".html",
                    ".css",
                    ".json",
                    ".xml",
                    ".yaml",
                    ".yml",
                    ".md",
                    ".txt",
                    ".csv",
                    ".sql",
                    ".sh",
                    ".bat",
                    ".ps1",
                    ".toml",
                    ".ini",
                    ".cfg",
                    ".jpg",
                    ".jpeg",
                    ".png",
                    ".gif",
                    ".svg",
                    ".pdf",
                    ".doc",
                    ".docx",
                }
                path = safe_path_join(
                    workspace_dir,
                    path,
                    allow_absolute=False,
                    allowed_extensions=allowed_extensions,
                )
            except (PathTraversalError, ValueError) as e:
                logger.exception("Path validation failed: ")
                return f"Error: Invalid file path - {e!s}"

        # Read existing file
        path_obj = Path(path)
        if path_obj.exists():
            content_lines = path_obj.read_text(encoding="utf-8").splitlines(keepends=True)
        else:
            # Create new file if it doesn't exist
            content_lines = []

        # Prepare content to insert
        insert_lines = content.split("\n")
        if not insert_lines[-1].endswith("\n"):
            insert_lines[-1] += "\n"

        # Insert content
        if line == 0:  # Append to end
            content_lines.extend(insert_lines)
        else:
            # Convert to 0-based index
            insert_idx = min(line - 1, len(content_lines))
            content_lines[insert_idx:insert_idx] = insert_lines

        # Write back to file
        dir_name = path_obj.parent
        if dir_name != Path():
            dir_name.mkdir(parents=True, exist_ok=True)
        path_obj.write_text("".join(content_lines), encoding="utf-8")

        action = "appended to" if line == 0 else f"inserted at line {line}"
        return f"Successfully inserted content {action} in {path_obj.name}"


# Write to File Tool
class WriteToFileInput(BaseModel):
    """Input for WriteToFileTool."""

    path: str = Field(..., description="File path to write to, relative to workspace directory.")
    content: str = Field(..., description="Complete content to write to the file.")
    line_count: int | None = Field(
        None,
        description="Total number of lines in the content (optional, will be calculated if not provided).",
    )


class WriteToFileTool(CustomBaseTool):
    """Tool for creating new files or completely rewriting existing files."""

    name: str = "write_to_file"
    description: str = (
        "Creates new files or completely rewrites existing files with the provided content. "
        "Supports plain text files and binary document formats (DOCX, DOC, PDF). "
        "For DOCX/DOC/PDF, content is interpreted as Markdown/plain text and converted to the target format. "
        "All file paths are relative to the current workspace directory."
    )
    args_schema: type[BaseModel] = WriteToFileInput

    # Extensions that need special binary document writing
    DOCUMENT_EXTENSIONS: set[str] = {".pdf", ".docx", ".doc"}

    @safe_tool_operation("write_to_file", fallback_value="Error: Failed to write file")
    def _run(self, path: str, content: str, line_count: int | None = None) -> str:
        """Write content to file."""
        # Get workspace directory from context
        workspace_dir = None
        if hasattr(self, "context") and self.context is not None:
            workspace_dir = getattr(self.context, "cwd", None)

        # Sanitize filename to prevent directory traversal
        sanitized_filename = sanitize_filename(Path(path).name)

        # Reconstruct path with sanitized filename
        path_obj = Path(path)
        dir_name = path_obj.parent
        path = str(dir_name / sanitized_filename) if dir_name != Path() else sanitized_filename

        # If we have a workspace directory, use safe path joining
        if workspace_dir:
            try:
                allowed_extensions: set[str] = {
                    ".py",
                    ".js",
                    ".ts",
                    ".tsx",
                    ".jsx",
                    ".vue",
                    ".html",
                    ".css",
                    ".json",
                    ".xml",
                    ".yaml",
                    ".yml",
                    ".md",
                    ".txt",
                    ".csv",
                    ".sql",
                    ".sh",
                    ".bat",
                    ".ps1",
                    ".toml",
                    ".ini",
                    ".cfg",
                    ".jpg",
                    ".jpeg",
                    ".png",
                    ".gif",
                    ".svg",
                    ".pdf",
                    ".doc",
                    ".docx",
                }
                path = safe_path_join(
                    workspace_dir,
                    path,
                    allow_absolute=False,
                    allowed_extensions=allowed_extensions,
                )
            except (PathTraversalError, ValueError) as e:
                logger.exception("Path validation failed: ")
                return f"Error: Invalid file path - {e!s}"

        # Create directory if it doesn't exist
        path_obj = Path(path)
        dir_name = path_obj.parent
        if dir_name != Path():
            dir_name.mkdir(parents=True, exist_ok=True)

        # Dispatch: binary documents vs plain text
        extension = path_obj.suffix.lower()
        if extension in self.DOCUMENT_EXTENSIONS:
            return self._write_document(path_obj, extension, content)

        return self._write_text_file(path_obj, content)

    def _write_text_file(self, path_obj: Path, content: str) -> str:
        """Write plain text file with UTF-8 encoding."""
        if content and not content.endswith("\n"):
            content += "\n"

        path_obj.write_text(content, encoding="utf-8")
        return f"File '{path_obj.name}' has been successfully written."

    def _write_document(self, path_obj: Path, extension: str, content: str) -> str:
        """Write binary document from text/Markdown content."""
        if extension == ".docx":
            return self._write_docx(path_obj, content)
        if extension == ".doc":
            return self._write_doc(path_obj, content)
        if extension == ".pdf":
            return self._write_pdf(path_obj, content)
        return f"Error: Unsupported document format: {extension}"

    @safe_tool_operation("write_docx", fallback_value="Error: Failed to write DOCX")
    def _write_docx(self, path_obj: Path, content: str) -> str:
        """Create a valid DOCX from Markdown. Tries pandoc first, falls back to python-docx."""
        if shutil.which("pandoc"):
            return self._write_docx_via_pandoc(path_obj, content)
        return self._write_docx_via_python(path_obj, content)

    # ------------------------------------------------------------------
    # Strategy 1: pandoc (best quality)
    # ------------------------------------------------------------------
    def _write_docx_via_pandoc(self, path_obj: Path, content: str) -> str:
        """Convert Markdown to DOCX via pandoc. Best format fidelity."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", encoding="utf-8", delete=False) as tmp:
            tmp.write(content)
            tmp_md = tmp.name

        try:
            proc = subprocess.run(
                ["pandoc", tmp_md, "-o", str(path_obj)],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if proc.returncode == 0:
                return f"File '{path_obj.name}' has been successfully written as DOCX (pandoc)."
            logger.warning("pandoc failed (rc=%d): %s", proc.returncode, proc.stderr)
            # Pandoc failed, fall back to python-docx
            return self._write_docx_via_python(path_obj, content)
        except subprocess.TimeoutExpired:
            logger.warning("pandoc timed out, falling back to python-docx")
            return self._write_docx_via_python(path_obj, content)
        finally:
            Path(tmp_md).unlink(missing_ok=True)

    # ------------------------------------------------------------------
    # Strategy 2: markdown-it-py + python-docx (pure Python fallback)
    # ------------------------------------------------------------------
    def _write_docx_via_python(self, path_obj: Path, content: str) -> str:
        """Convert Markdown to DOCX using markdown-it-py AST + python-docx."""
        try:
            from markdown_it import MarkdownIt
        except ImportError:
            return "Error: Neither pandoc nor markdown-it-py available. Install pandoc (apt install pandoc) or pip install markdown-it-py python-docx"
        try:
            import docx
        except ImportError:
            return "Error: python-docx not installed. Install with: pip install python-docx"

        md = MarkdownIt("commonmark", {"html": False}).enable("table")
        tokens = md.parse(content)

        doc = docx.Document()

        # Set default Chinese-friendly font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = docx.shared.Pt(11)
        try:
            style.element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "宋体")
        except Exception:
            pass

        # State for inline formatting within a paragraph
        _heading_level = 0
        _list_type: str | None = None  # "bullet" or "ordered"
        _list_depth = 0
        _list_number = 0
        _in_table = False
        _table_rows: list[list[str]] = []
        _table_cols = 0

        i = 0
        while i < len(tokens):
            tok = tokens[i]

            # --- Headings ---
            if tok.type == "heading_open":
                _heading_level = int(tok.tag[1])  # h1 -> 1, h2 -> 2, ...
                i += 1
                continue
            if tok.type == "heading_close":
                _heading_level = 0
                i += 1
                continue
            if tok.type == "inline" and _heading_level > 0:
                text = self._render_inline_text(tok)
                level = min(_heading_level, 4)
                doc.add_heading(text, level=level)
                _heading_level = 0
                i += 1
                continue

            # --- Paragraphs ---
            if tok.type == "paragraph_open":
                i += 1
                # Collect inline content
                if i < len(tokens) and tokens[i].type == "inline":
                    text = self._render_inline_text(tokens[i])
                    if text.strip():
                        doc.add_paragraph(text)
                    i += 1
                # skip paragraph_close
                if i < len(tokens) and tokens[i].type == "paragraph_close":
                    i += 1
                continue

            # --- Bullet lists ---
            if tok.type == "bullet_list_open":
                _list_type = "bullet"
                _list_depth += 1
                i += 1
                continue
            if tok.type == "bullet_list_close":
                _list_type = None
                _list_depth -= 1
                i += 1
                continue

            # --- Ordered lists ---
            if tok.type == "ordered_list_open":
                _list_type = "ordered"
                _list_number = int(tok.attrGet("start") or 1)
                _list_depth += 1
                i += 1
                continue
            if tok.type == "ordered_list_close":
                _list_type = None
                _list_depth = 0
                _list_number = 0
                i += 1
                continue

            # --- List items ---
            if tok.type == "list_item_open":
                i += 1
                # Collect inline content inside the list item
                text_parts: list[str] = []
                while i < len(tokens) and tokens[i].type not in (
                    "list_item_close",
                    "bullet_list_close",
                    "ordered_list_close",
                ):
                    if tokens[i].type == "inline":
                        text_parts.append(self._render_inline_text(tokens[i]))
                    i += 1
                text = "".join(text_parts)
                style_name = "List Bullet" if _list_type == "bullet" else "List Number"
                doc.add_paragraph(text, style=style_name)
                # skip list_item_close
                if i < len(tokens) and tokens[i].type == "list_item_close":
                    i += 1
                continue

            # --- Tables ---
            if tok.type == "table_open":
                _in_table = True
                _table_rows = []
                _table_cols = 0
                i += 1
                continue
            if tok.type == "table_close":
                _in_table = False
                # Write collected table rows to DOCX
                if _table_rows:
                    num_cols = max(len(row) for row in _table_rows) if _table_rows else 0
                    table = doc.add_table(rows=len(_table_rows), cols=num_cols)
                    table.style = "Table Grid"
                    for r_idx, row in enumerate(_table_rows):
                        for c_idx, cell_text in enumerate(row):
                            if c_idx < num_cols:
                                table.cell(r_idx, c_idx).text = cell_text
                _table_rows = []
                i += 1
                continue
            if tok.type == "tr_open":
                i += 1
                current_row: list[str] = []
                while i < len(tokens) and tokens[i].type != "tr_close":
                    if tokens[i].type == "inline":
                        current_row.append(self._render_inline_text(tokens[i]).strip())
                    i += 1
                _table_rows.append(current_row)
                if i < len(tokens) and tokens[i].type == "tr_close":
                    i += 1
                continue

            # --- Code blocks ---
            if tok.type in ("code_block", "fence"):
                code_text = tok.content.rstrip()
                if code_text:
                    para = doc.add_paragraph(code_text)
                    para.style = doc.styles["Normal"]
                    for run in para.runs:
                        run.font.name = "Courier New"
                        run.font.size = docx.shared.Pt(9)
                i += 1
                continue

            # --- Horizontal rules ---
            if tok.type == "hr":
                doc.add_paragraph("—")
                i += 1
                continue

            # --- Skip everything else ---
            i += 1

        doc.save(str(path_obj))
        return f"File '{path_obj.name}' has been successfully written as DOCX (python-docx)."

    @staticmethod
    def _render_inline_text(tok) -> str:
        """Render markdown-it inline token children to plain text."""
        parts: list[str] = []
        if not hasattr(tok, "children") or not tok.children:
            return tok.content if tok.content else ""

        for child in tok.children:
            if child.type == "text":
                parts.append(child.content)
            elif child.type in ("softbreak", "hardbreak"):
                parts.append("\n")
            elif child.type in ("strong_open", "strong_close", "em_open", "em_close"):
                pass  # Strip formatting markers for plain text output
            elif child.type in ("link_open", "link_close"):
                if child.type == "link_open":
                    href = child.attrGet("href") or ""
                    if href:
                        parts.append("[")
                else:
                    href = ""
                    # Find the matching link_open to get href
                    parts.append("]")
            elif child.type == "code_inline":
                parts.append(f"`{child.content}`")
            else:
                parts.append(child.content if child.content else "")
        return "".join(parts)

    # ------------------------------------------------------------------
    # DOC: write DOCX first, then convert via libreoffice
    # ------------------------------------------------------------------
    @safe_tool_operation("write_doc", fallback_value="Error: Failed to write DOC")
    def _write_doc(self, path_obj: Path, content: str) -> str:
        """Create a .doc file by writing DOCX first then converting via libreoffice."""
        tmp_docx = path_obj.with_suffix(".tmp.docx")
        result = self._write_docx(tmp_docx, content)
        if result.startswith("Error:"):
            tmp_docx.unlink(missing_ok=True)
            return result

        try:
            proc = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "doc",
                    "--outdir",
                    str(path_obj.parent),
                    str(tmp_docx),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            tmp_docx.unlink(missing_ok=True)

            if proc.returncode == 0 and path_obj.exists():
                return f"File '{path_obj.name}' has been successfully written as DOC."
            return f"Error: libreoffice conversion failed: {proc.stderr.strip()}"
        except FileNotFoundError:
            tmp_docx.unlink(missing_ok=True)
            return "Error: libreoffice not installed. Cannot create .doc files without libreoffice."
        except subprocess.TimeoutExpired:
            tmp_docx.unlink(missing_ok=True)
            return "Error: libreoffice conversion timed out."

    # ------------------------------------------------------------------
    # PDF: pandoc (best) → fpdf2 (fallback)
    # ------------------------------------------------------------------
    @safe_tool_operation("write_pdf", fallback_value="Error: Failed to write PDF")
    def _write_pdf(self, path_obj: Path, content: str) -> str:
        """Create a PDF from Markdown. Tries pandoc first, falls back to fpdf2."""
        if shutil.which("pandoc"):
            return self._write_pdf_via_pandoc(path_obj, content)
        return self._write_pdf_via_fpdf(path_obj, content)

    def _write_pdf_via_pandoc(self, path_obj: Path, content: str) -> str:
        """Convert Markdown to PDF via pandoc."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", encoding="utf-8", delete=False) as tmp:
            tmp.write(content)
            tmp_md = tmp.name

        try:
            proc = subprocess.run(
                ["pandoc", tmp_md, "-o", str(path_obj)],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if proc.returncode == 0:
                return f"File '{path_obj.name}' has been successfully written as PDF (pandoc)."
            logger.warning("pandoc pdf failed (rc=%d): %s", proc.returncode, proc.stderr)
            return self._write_pdf_via_fpdf(path_obj, content)
        except subprocess.TimeoutExpired:
            return self._write_pdf_via_fpdf(path_obj, content)
        finally:
            Path(tmp_md).unlink(missing_ok=True)

    @safe_tool_operation("write_pdf_fpdf", fallback_value="Error: Failed to write PDF")
    def _write_pdf_via_fpdf(self, path_obj: Path, content: str) -> str:
        """Create a PDF from Markdown text using fpdf2 (pure Python fallback)."""
        try:
            from fpdf import FPDF
        except ImportError:
            return "Error: Neither pandoc nor fpdf2 available. Install pandoc or pip install fpdf2"

        pdf = FPDF()
        pdf.add_page()

        # Try to register a CJK font
        cjk_font = None
        for font_name, font_path in [
            ("WQY", "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            ("Noto", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
            ("ARPL", "/usr/share/fonts/truetype/arphic/uming.ttc"),
        ]:
            if Path(font_path).exists():
                pdf.add_font(font_name, "", font_path, uni=True)
                cjk_font = font_name
                break

        def set_font(size: int = 11) -> None:
            if cjk_font:
                pdf.set_font(cjk_font, size=size)
            else:
                pdf.set_font("Helvetica", size=size)

        set_font()
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                pdf.ln(5)
                continue

            # Count heading level
            heading_match = re.match(r"^(#{1,4})\s", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = stripped[heading_match.end() :]
                font_size = max(11, 18 - level * 2)
                set_font(font_size)
                pdf.multi_cell(0, 8, text)
                pdf.ln(2)
            else:
                set_font()
                pdf.multi_cell(0, 6, stripped)

        pdf.output(str(path_obj))
        return f"File '{path_obj.name}' has been successfully written as PDF (fpdf2)."


# Enhanced Apply Diff Tool (improved version)
class ApplyDiffInput(BaseModel):
    """Input for ApplyDiffTool."""

    file_path: str = Field(..., description="Path to file to apply the diff.")
    diff: str = Field(
        ...,
        description="The diff content in either SEARCH/REPLACE format or git unified diff format. Both are auto-detected and supported.",
    )
    verify: bool = Field(
        False,
        description="Verify the diff was applied correctly by checking the file content.",
    )


class ApplyDiffTool(CustomBaseTool):
    """Tool for applying precise diff patches to files."""

    name: str = "apply_diff"
    description: str = """Makes precise, surgical changes to files using diff format.

**IMPORTANT - For Large Files:**
When editing files larger than 100 lines, split changes into multiple smaller diffs:
- Each diff should modify one specific section (10-50 lines)
- Apply diffs sequentially, waiting for confirmation before next
- This prevents response truncation and ensures accuracy

**Supported Formats (auto-detected):**

**1. Git Unified Diff Format (Recommended for compatibility):**
```diff
--- a/file.txt
+++ b/file.txt
@@ -line,count +line,count @@
 context line
-old line to remove
+new line to add
 another context
```

**2. SEARCH/REPLACE Format (Simpler for LLMs):**
```
<<<<<<< SEARCH
[exact content to find]
=======
[replacement content]
>>>>>>> REPLACE
```

**Best Practices:**
- Include enough context for unique matching (3-5 lines)
- Keep changes focused and minimal
- For large HTML/code files, use apply_diff instead of write_to_file"""
    args_schema: type[BaseModel] = ApplyDiffInput

    @safe_tool_operation("apply_diff", fallback_value="Error: Failed to apply diff")
    def _run(self, file_path: str, diff: str, verify: bool = False) -> str:
        """Apply diff patch to file."""
        # Get workspace directory from context
        workspace_dir = None
        if hasattr(self, "context") and self.context is not None:
            workspace_dir = getattr(self.context, "cwd", None)

        # Sanitize filename to prevent directory traversal
        sanitized_filename = sanitize_filename(Path(file_path).name)

        # Reconstruct path with sanitized filename
        path_obj = Path(file_path)
        dir_name = path_obj.parent
        file_path = str(dir_name / sanitized_filename) if dir_name != Path() else sanitized_filename

        # If we have a workspace directory, use safe path joining
        if workspace_dir:
            try:
                # Allow common file extensions
                allowed_extensions: set[str] = {
                    ".py",
                    ".js",
                    ".ts",
                    ".tsx",
                    ".jsx",
                    ".vue",
                    ".html",
                    ".css",
                    ".json",
                    ".xml",
                    ".yaml",
                    ".yml",
                    ".md",
                    ".txt",
                    ".csv",
                    ".sql",
                    ".sh",
                    ".bat",
                    ".ps1",
                    ".toml",
                    ".ini",
                    ".cfg",
                    ".jpg",
                    ".jpeg",
                    ".png",
                    ".gif",
                    ".svg",
                    ".pdf",
                    ".doc",
                    ".docx",
                }
                file_path = safe_path_join(
                    workspace_dir,
                    file_path,
                    allow_absolute=False,
                    allowed_extensions=allowed_extensions,
                )
            except (PathTraversalError, ValueError) as e:
                logger.exception("Path validation failed: ")
                return f"Error: Invalid file path - {e!s}"

        path_obj = Path(file_path)
        if not path_obj.exists():
            return f"Error: File not found at {path_obj.name}"

        # Read original file
        original_content = path_obj.read_text(encoding="utf-8")
        original_lines = original_content.count("\n") + 1

        # Parse diff blocks
        diff_blocks = self._parse_diff(diff)

        if not diff_blocks:
            return "Error: No valid diff blocks found. Make sure your diff follows either SEARCH/REPLACE or git unified diff format."

        # Apply each diff block
        modified_content = original_content
        success_count = 0
        failed_blocks = []

        for i, block in enumerate(diff_blocks, 1):
            search_snippet = block["search"][:100]
            result = self._apply_diff_block(modified_content, block)

            if result is None:
                failed_blocks.append(
                    {
                        "block_num": i,
                        "search_preview": search_snippet,
                        "line_number": self._find_line_number(
                            modified_content,
                            block["search"][:50],
                        ),
                    },
                )
            else:
                modified_content = result
                success_count += 1

        # Check if all blocks succeeded
        if failed_blocks:
            error_msg = f"Applied {success_count}/{len(diff_blocks)} diffs successfully. Failed blocks:\n"
            for failed in failed_blocks:
                error_msg += f"\n  Block {failed['block_num']}: Search content not found near line {failed['line_number']}"
                error_msg += f"\n    Preview: {failed['search_preview']}"
            return error_msg

        # Write modified content back
        path_obj.write_text(modified_content, encoding="utf-8")

        # Calculate stats
        new_lines = modified_content.count("\n") + 1
        lines_added = max(0, new_lines - original_lines)
        lines_removed = max(0, original_lines - new_lines)

        result_msg = f"Successfully applied {success_count} diff block(s) to {path_obj.name}"
        if lines_added or lines_removed:
            result_msg += f" (Lines: +{lines_added}, -{lines_removed})"

        # Verify if requested
        if verify:
            verification = self._verify_diff(file_path, diff_blocks)
            result_msg += f"\n\nVerification: {verification}"

        return result_msg

    def _find_line_number(self, content: str, search_snippet: str) -> int:
        """Find approximate line number for a search snippet."""
        lines = content.split("\n")
        snippet_lower = search_snippet.lower()[:50]

        for i, line in enumerate(lines, 1):
            if snippet_lower in line.lower():
                return i

        return -1

    def _verify_diff(self, file_path: str, diff_blocks: list[dict]) -> str:
        """Verify that all diffs were applied correctly."""
        try:
            path_obj = Path(file_path)
            content = path_obj.read_text(encoding="utf-8")

            verified_count = 0
            for block in diff_blocks:
                # Check if replacement text exists in file
                if block["replace"].rstrip("\n") in content:
                    verified_count += 1

            if verified_count == len(diff_blocks):
                return f"All {verified_count} changes verified successfully"
            return f"{verified_count}/{len(diff_blocks)} changes verified"

        except FileNotFoundError:
            return f"Verification failed: File not found at {file_path}"
        except OSError as e:
            return f"Verification failed: Unable to read file - {e!s}"
        except (KeyError, AttributeError) as e:
            # Diff block structure error
            return f"Verification failed: Invalid diff block structure - {e!s}"

    def _parse_diff(self, diff: str) -> list[dict]:
        """Parse diff content into blocks. Supports both SEARCH/REPLACE and git unified diff formats."""
        # Auto-detect format
        if "<<<<<<< SEARCH" in diff or "<<<<<<< SEARCH" in diff:
            return self._parse_search_replace(diff)
        if re.search(r"^--- a/", diff, re.MULTILINE) or re.search(
            r"^@@ -\d+,\d+ \+\d+,\d+ @@",
            diff,
            re.MULTILINE,
        ):
            return self._parse_unified_diff(diff)
        # Try unified diff first as fallback
        logger.warning("Unable to detect diff format, trying unified diff parser")
        return self._parse_unified_diff(diff)

    def _parse_search_replace(self, diff: str) -> list[dict]:
        """Parse SEARCH/REPLACE format diff."""
        blocks = []
        current_block = None

        lines = diff.split("\n")
        for line in lines:
            if line.startswith("<<<<<<< SEARCH"):
                if current_block:
                    blocks.append(current_block)
                current_block = {"search": "", "replace": ""}
            elif line.startswith("======="):
                if current_block:
                    current_block["mode"] = "replace"
            elif line.startswith(">>>>>>> REPLACE"):
                if current_block:
                    blocks.append(current_block)
                current_block = None
            elif current_block:
                if "mode" not in current_block:
                    current_block["search"] += line + "\n"
                else:
                    current_block["replace"] += line + "\n"

        return blocks

    def _parse_unified_diff(self, diff: str) -> list[dict]:
        """Parse git unified diff format."""
        blocks = []
        lines = diff.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Look for hunk header @@ -old_start,old_count +new_start,new_count @@
            hunk_match = re.match(r"^@@ -(\d+),?(\d+)? \+(\d+),?(\d+)? @@", line)
            if hunk_match:
                int(hunk_match.group(1))
                int(hunk_match.group(2) or 1)
                int(hunk_match.group(3))
                int(hunk_match.group(4) or 1)

                # Collect hunk content
                context_lines = []
                old_lines = []
                new_lines = []
                j = i + 1

                while j < len(lines):
                    hunk_line = lines[j]

                    # Stop at next hunk or file header
                    if hunk_line.startswith(("@@", "---", "+++")):
                        break

                    # Parse hunk lines
                    if hunk_line.startswith(" ") or not hunk_line:
                        # Context line
                        context_lines.append(
                            hunk_line.removeprefix(" "),
                        )
                        old_lines.append(hunk_line.removeprefix(" "))
                        new_lines.append(hunk_line.removeprefix(" "))
                    elif hunk_line.startswith("-"):
                        # Removed line
                        old_lines.append(hunk_line[1:])
                    elif hunk_line.startswith("+"):
                        # Added line
                        new_lines.append(hunk_line[1:])

                    j += 1

                # Create search/replace block from hunk
                search_text = "\n".join(old_lines).rstrip("\n")
                replace_text = "\n".join(new_lines).rstrip("\n")

                if search_text or replace_text:
                    blocks.append(
                        {
                            "search": search_text + "\n" if search_text else "",
                            "replace": replace_text + "\n" if replace_text else "",
                        },
                    )

                i = j - 1  # Will increment to j in loop

            i += 1

        return blocks

    def _apply_diff_block(self, content: str, block: dict) -> str | None:
        """Apply a single diff block to content."""
        search_text = block["search"].rstrip("\n")
        replace_text = block["replace"].rstrip("\n")

        # First try exact match
        if search_text in content:
            return content.replace(search_text, replace_text, 1)

        # Try with trailing newline variations
        variations = [
            search_text + "\n",
            search_text + "\r\n",
            search_text.rstrip("\n"),
        ]

        for variation in variations:
            if variation in content:
                return content.replace(variation, replace_text, 1)

        # If not found, provide helpful context
        lines = content.split("\n")
        search_lines = search_text.split("\n")

        # Try to find partial match for context
        for i in range(len(lines) - len(search_lines) + 1):
            window = "\n".join(lines[i : i + len(search_lines)])
            similarity = self._calculate_similarity(search_text[:100], window[:100])

            if similarity > 0.5:
                return None  # Indicate potential match with context

        return None

    def _calculate_similarity(self, str1: str, str2: str) -> float:
        """Calculate simple similarity ratio between two strings."""
        if not str1 or not str2:
            return 0.0

        set1 = set(str1.lower().split())
        set2 = set(str2.lower().split())

        if not set1 or not set2:
            return 0.0

        intersection = set1.intersection(set2)
        union = set1.union(set2)

        return len(intersection) / len(union) if union else 0.0
